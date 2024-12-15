import io
import os
import subprocess
import tempfile

import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QFileDialog, QSizeGrip
from PyQt5.uic import loadUi

from configuration.configuration_buttom import (
    icon_configurate_manager,
    icon_configurate_top,
    icon_excel,
    icon_exit_program,
)
from configuration.configuration_buttom_top import (
    control_bt_maximizar,
    control_bt_minimizar,
    control_bt_normal,
)
from configuration.configuration_config_theme import load_config
from configuration.configuration_delete_banner import delete_banner
from configuration.configuration_window_move import mousePressEvent, window_move
from controller.controllerbusiness import BusinessController
from controller.controllerreport import ReportController
from view.admin.report.viewadd import Viewadd
from view.admin.task.viewupdate import Viewupdate


class Viewmainreport(QtWidgets.QMainWindow):
    def __init__(self):
        super(Viewmainreport, self).__init__()
        self.theme = load_config(self)  # Lee la configuración al iniciar
        loadUi(f"design/standar/mainreport{self.theme}.ui", self)

        icon_configurate_top(self)
        icon_exit_program(self)
        icon_excel(self)
        delete_banner(self)
        # Inicializar el QSizeGrip y establecer su tamaño
        self.gripSize = 16  # Define el tamaño del grip
        self.grip = QSizeGrip(self)  # Crea un QSizeGrip para el redimensionamiento
        self.grip.resize(self.gripSize, self.gripSize)  # Ajusta el tamaño del QSizeGrip

        # Maximizar la ventana por defecto
        self.showMaximized()

        # Asigna los eventos personalizados al frame superior
        self.frame_superior.mousePressEvent = lambda event: mousePressEvent(self, event)
        self.frame_superior.mouseMoveEvent = lambda event: window_move(self, event)

        # Control de botones de la barra de títulos
        self.bt_minimizar.clicked.connect(lambda: control_bt_minimizar(self))
        self.bt_restaurar.clicked.connect(lambda: control_bt_normal(self))
        self.bt_maximizar.clicked.connect(lambda: control_bt_maximizar(self))
        self.bt_cerrar.clicked.connect(lambda: self.close())
        self.bt_maximizar.hide()

        self.btn_get.clicked.connect(self.show)

        self.btn_excel.clicked.connect(self.export_excel)
        self.btn_search.clicked.connect(self.search)
        self.btn_exit.clicked.connect(self.close_program)
        self.btn_report_export.clicked.connect(self.export_word)

        self.controller = ReportController(self)
        self.controllerbusiness = BusinessController(self)

        self.show()

    def export_word(self):
        try:
            rows = self.controllerbusiness.get()

            if not rows:
                QtWidgets.QMessageBox.warning(
                    self, "Advertencia", "No se encontraron datos para exportar."
                )
                return

            # Crear un nuevo documento de Word
            doc = Document()

            # Crear la portada del documento
            doc.add_heading("Informe Empresarial", 0)

            # Asumiendo que hay al menos una fila
            first_row = rows[0]
            doc.add_heading("Datos de la Empresa", level=1)
            doc.add_paragraph(f"ID: {first_row[0]}")
            doc.add_paragraph(f"Nombre del Negocio: {first_row[1]}")
            doc.add_paragraph(f"Número legal: {first_row[3]}")
            doc.add_paragraph(f"Industria: {first_row[4]}")
            doc.add_paragraph(f"Número de registro: {first_row[5]}")
            doc.add_paragraph(f"Fecha de fundación: {first_row[6]}")

            # Manejar el BLOB de imagen
            image_blob = first_row[
                2
            ]  # Asumiendo que row[2] contiene el BLOB de la imagen
            if image_blob:
                image = Image.open(io.BytesIO(image_blob))
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=".png"
                ) as temp_image:
                    image.save(temp_image, format="PNG")
                    image_path = temp_image.name
                doc.add_paragraph().add_run().add_picture(image_path, width=Inches(2))
                os.remove(image_path)  # Eliminar el archivo temporal

            doc.add_page_break()

            # Agregar una tabla al documento con el número adecuado de filas y columnas
            table = doc.add_table(rows=1, cols=7)
            # Definir los encabezados de la tabla
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ID"
            hdr_cells[1].text = "Nombre del Negocio"
            hdr_cells[2].text = "Imagen"
            hdr_cells[3].text = "Número legal"
            hdr_cells[4].text = "Industria"
            hdr_cells[5].text = "Número de registro"
            hdr_cells[6].text = "Fecha de fundación"

            # Insertar los datos de la base de datos en la tabla
            for row in rows:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row[0])  # ID
                row_cells[1].text = row[1]  # Nombre del Negocio
                row_cells[3].text = row[3]  # Número legal
                row_cells[4].text = row[4]  # Industria
                row_cells[5].text = row[5]  # Número de registro
                row_cells[6].text = row[6]  # Fecha de fundación

                # Manejar el BLOB de imagen
                image_blob = row[
                    2
                ]  # Asumiendo que row[2] contiene el BLOB de la imagen
                if image_blob:
                    image = Image.open(io.BytesIO(image_blob))
                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=".png"
                    ) as temp_image:
                        image.save(temp_image, format="PNG")
                        image_path = temp_image.name
                    row_cells[2].add_paragraph().add_run().add_picture(
                        image_path, width=Inches(1.25)
                    )
                    os.remove(image_path)  # Eliminar el archivo temporal

            # Guardar el documento de Word

            file_path = os.path.join(
                os.path.dirname(__file__), "../../../doc/Informe_empresarial.docx"
            )
            doc.save(file_path)

            # Mostrar mensaje de confirmación
            QtWidgets.QMessageBox.information(
                self, "Exportar a Word", f"El informe ha sido exportado a {file_path}"
            )

            # Abrir la ubicación del archivo generado
            folder_path = os.path.dirname(os.path.abspath(file_path))
            if os.name == "nt":  # Windows
                subprocess.Popen(f"explorer /select,{os.path.abspath(file_path)}")
            else:
                # Linux y macOS
                subprocess.Popen(["xdg-open", folder_path])

        except Exception as e:
            QtWidgets.QMessageBox.critical(
                self, "Error", f"Se produjo un error al exportar el documento: {str(e)}"
            )

    def close_program(self):
        QApplication.quit()

    def show(self):
        # Obtener datos desde el controlador (usando el INNER JOIN)
        Report_controller = self.controller.get()

        # Limpiar la tabla antes de insertar nuevos datos
        self.table_report.setRowCount(0)

        # Configuración del tamaño de las filas y columnas
        self.table_report.verticalHeader().setDefaultSectionSize(200)
        self.table_report.horizontalHeader().setDefaultSectionSize(300)

        # Mostrar los datos obtenidos en la tabla
        for i, report in enumerate(Report_controller):
            self.table_report.insertRow(i)

            # Verifica si `task` es una tupla o lista con múltiples elementos
            if isinstance(report, (tuple, list)):
                # Asumiendo que la consulta INNER JOIN devuelve al menos estos campos:
                # (user_id, user_name, task_id, task_description, task_status)
                self.table_report.setItem(
                    i, 0, QtWidgets.QTableWidgetItem(str(report[0]))
                )  # user_id
                self.table_report.setItem(
                    i, 1, QtWidgets.QTableWidgetItem(report[1])
                )  # user_name
                self.table_report.setItem(
                    i, 2, QtWidgets.QTableWidgetItem(report[2])
                )  # task_id
                self.table_report.setItem(
                    i, 3, QtWidgets.QTableWidgetItem(report[3])
                )  # task_name
                self.table_report.setItem(
                    i, 4, QtWidgets.QTableWidgetItem(report[4])
                )  # task_description
                self.table_report.setItem(
                    i, 5, QtWidgets.QTableWidgetItem(report[5])
                )  # task_status

                self.table_report.setItem(
                    i, 6, QtWidgets.QTableWidgetItem(report[6])
                )  # task_status

            else:
                # Si `task` no es una tupla o lista, maneja el error
                print(
                    f"Error: Se esperaba una tupla, pero se encontró un valor único: {report}"
                )

    def export_excel(self):
        # Obtener datos de usuarios desde el controlador
        user_controller = self.controller.get_user()

        # Convertir los datos a un DataFrame de pandas
        df = pd.DataFrame(
            user_controller,
            columns=[
                "ID",
                "Nombre",
                "Correo",
                "Tel. Contacto",
                "Dirección",
                "Contraseña",
                "Rol del usuario",
            ],
        )

        # Mostrar el cuadro de diálogo para guardar el archivo
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar Usuario",
            "",
            "Archivos de Excel (*.xlsx);;Todos los archivos (*)",
            options=options,
        )

        if file_path:
            # Guardar el DataFrame en un archivo de Excel
            df.to_excel(file_path, index=False)

            # Cargar el archivo Excel para aplicar estilos
            wb = load_workbook(file_path)
            ws = wb.active

            # Aplicar estilos a los encabezados
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill("solid", fgColor="4F81BD")
            header_alignment = Alignment(horizontal="center", vertical="center")

            for cell in ws[1]:  # Primera fila contiene los encabezados
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment

            # Guardar los cambios en el archivo
            wb.save(file_path)

            QtWidgets.QMessageBox.information(
                self, "Éxito", f"Usuario exportado correctamente a {file_path}"
            )
        else:
            QtWidgets.QMessageBox.warning(
                self, "Cancelado", "La exportación fue cancelada."
            )

    def search(self):
        search_text = self.searchtxt.text().lower()  # Texto de búsqueda en minúsculas
        found = False  # Variable para rastrear si se encontró una coincidencia

        for row in range(self.table_client.rowCount()):
            match = False
            for col in range(self.table_client.columnCount()):
                item = self.table_client.item(row, col)
                if item and search_text in item.text().lower():  # Comparar texto
                    match = True
                    found = True  # Si hay una coincidencia, actualizamos 'found'
                    break  # Si encuentra coincidencia, no sigue buscando en esa fila

            # Ocultar la fila si no coincide
            self.table_client.setRowHidden(row, not match)

        # Si no se encontró ninguna coincidencia, mostrar un mensaje de advertencia
        if not found:
            QtWidgets.QMessageBox.warning(
                self,
                "Sin coincidencias",
                "No se encontraron resultados para la búsqueda.",
            )
