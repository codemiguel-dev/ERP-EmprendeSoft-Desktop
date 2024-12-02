# configuraciones_gui.py
import os

from docx import Document


def report_doc(id_user, name, type_report, description, file_name):
    # Crear el documento de Word
    doc = Document()
    doc.add_heading(type_report, 0)
    doc.add_paragraph(f"Nombre: {name}")
    doc.add_paragraph(f"Descripción: {description}")
    doc.add_paragraph(f"Tipo de informe: {type_report}")
    doc.add_paragraph(f"ID de Usuario: {id_user}")

    # Especificar la carpeta y el nombre del archivo
    directory = "doc"
    if not os.path.exists(directory):
        os.makedirs(directory)  # Crear la carpeta si no existe

    file_path = os.path.join(directory, file_name)

    # Guardar el documento
    doc.save(file_path)
